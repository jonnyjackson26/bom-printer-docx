from docx import Document
import os
import subprocess
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    books = ["title", "introduction", "three", "eight", "js", "1-nephi"]
    chapters = {
        "title": 1, "introduction": 1, "three": 1, "eight": 1, "js": 1,
        "1-nephi": 1
    }

    document = Document()

    # Set smaller margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(.5)
        section.bottom_margin = Inches(.5)
        section.left_margin = Inches(.4)
        section.right_margin = Inches(.4)
        section.footer_distance = Inches(0.2)

    def style_cell_text(cell, text, font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, italic=False):
        cell.text = ''
        run = cell.paragraphs[0].add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = bold
        run.italic = italic
        cell.paragraphs[0].alignment = alignment
        paragraph_format = cell.paragraphs[0].paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(4)
        paragraph_format.line_spacing = Pt(12)

        # Apply cell borders
        tc_pr = cell._element.get_or_add_tcPr()
        borders = tc_pr.find(qn('w:tcBorders'))
        if borders is None:
            borders = OxmlElement('w:tcBorders')
            tc_pr.append(borders)
        
        for border in ['top', 'left', 'bottom', 'right']:
            b = OxmlElement(f'w:{border}')
            b.set(qn('w:val'), 'nil')
            b.set(qn('w:space'), '0')
            borders.append(b)

    def add_horizontal_line(doc):
        p = doc.add_paragraph()
        run = p.add_run()
        hr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        hr.append(bottom)
        p._element.get_or_add_pPr().append(hr)

    def add_title_page(doc):
        doc.add_paragraph("\n\n\n")
        main_title = doc.add_paragraph()
        main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = main_title.add_run("The Book of Mormon")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(36)
        run.bold = True
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Another testament of Jesus Christ")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(24)
        run.italic = True
        doc.add_page_break()

    add_title_page(document)

    num_lines_per_column = 40  # Adjust this number based on your column height

    for book in books:
        add_horizontal_line(document)
        document.add_paragraph("")
        for chapter in range(1, chapters[book] + 1):
            path = f'bom-english/{book}/{chapter}.txt'
            with open(path, 'r', encoding='utf-8') as file:
                verses = [line.strip() for line in file.readlines() if line.strip()]
            
            # Create a table with two columns
            table = document.add_table(rows=0, cols=2)
            table.autofit = False
            table.columns[0].width = Inches(3)  # Adjust width as needed
            table.columns[1].width = Inches(3)  # Adjust width as needed

            # Center the table
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add text to columns
            current_line = 0
            while current_line < len(verses):
                left_text = "\n".join(verses[current_line:current_line + num_lines_per_column])
                current_line += num_lines_per_column
                right_text = "\n".join(verses[current_line:current_line + num_lines_per_column])
                current_line += num_lines_per_column

                row_cells = table.add_row().cells
                style_cell_text(row_cells[0], left_text)
                style_cell_text(row_cells[1], right_text)

                # Add vertical line between columns
                cell_left = row_cells[0]
                tc_pr_left = cell_left._element.get_or_add_tcPr()
                borders_left = tc_pr_left.find(qn('w:tcBorders'))
                if borders_left is None:
                    borders_left = OxmlElement('w:tcBorders')
                    tc_pr_left.append(borders_left)
                right_border = OxmlElement('w:right')
                right_border.set(qn('w:val'), 'single')
                right_border.set(qn('w:sz'), '4')
                right_border.set(qn('w:space'), '0')
                borders_left.append(right_border)

                cell_right = row_cells[1]
                tc_pr_right = cell_right._element.get_or_add_tcPr()
                borders_right = tc_pr_right.find(qn('w:tcBorders'))
                if borders_right is None:
                    borders_right = OxmlElement('w:tcBorders')
                    tc_pr_right.append(borders_right)
                left_border = OxmlElement('w:left')
                left_border.set(qn('w:val'), 'single')
                left_border.set(qn('w:sz'), '4')
                left_border.set(qn('w:space'), '0')
                borders_right.append(left_border)

            add_horizontal_line(document)

    def add_page_numbers(document):
        sections = document.sections
        for section in sections:
            footer = section.footer
            p = footer.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            field = OxmlElement('w:fldSimple')
            field.set(qn('w:instr'), 'PAGE')
            run._element.append(field)

    add_page_numbers(document)
    document.save("bom.docx")
    subprocess.call(['powershell.exe', 'Start-Process', 'bom.docx'])

main()
