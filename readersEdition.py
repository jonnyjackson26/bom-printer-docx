from docx import Document
import os
import subprocess
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from info import books_dict

def page_formatting(document):
    sections = document.sections
    for section in sections:
        section.page_width = Inches(6)
        section.page_height = Inches(9)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.footer_distance = Inches(0.5)

def add_horizontal_line(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    hr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')  # Thinner line
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), '000000')  # Black color
    hr.append(bottom)
    p._element.get_or_add_pPr().append(hr)
def add_page_break(document):
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break()  # Adds a page break to the document

def add_content(document):
    for book_key in books_dict:
        book = books_dict[book_key]
        addParagraph(f"{book['name']}","book-title",document)
        add_page_break(document)
        document.add_paragraph("")
        for chapter in range(1, book["numOfChapters"] + 1):
            path = f'bom-english/{book_key}/{chapter}.txt'
            with open(path, 'r', encoding='utf-8') as file:
                verses = [line.strip() for line in file.readlines() if line.strip()]

             # Combine verses into paragraphs of 3 verses each
            combined_verses = [ ' '.join(verses[i:i+3]) for i in range(0, len(verses), 3)]
            
            if(book["numOfChapters"]>1):
                addParagraph(f"{book['name']} {chapter}","chapter-title",document) #Chapter 1
            else:
                addParagraph(f"{book['name']}","chapter-title",document)
            for paragraph in combined_verses:
                addParagraph(f"{paragraph}","normal",document)

            add_page_break(document)

def addParagraph(text,mode,document):
    if mode=="normal":
        p = document.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Cambria'  
        run.font.size = Pt(12)  
        p.paragraph_format.first_line_indent = Pt(24)
        p.paragraph_format.line_spacing = Pt(14)  # Add line spacing for readability
    if mode=="chapter-title":
        p = document.add_paragraph()
        run = p.add_run(text.upper())
        run.font.name = 'Georgia'  
        run.font.size = Pt(13)  
        p.paragraph_format.first_line_indent = Pt(24)
        p.paragraph_format.line_spacing = Pt(14)
        p.alignment= WD_ALIGN_PARAGRAPH.CENTER
    if mode=="book-title":
        p = document.add_paragraph()
        run = p.add_run(text.upper())
        run.font.name = 'Lora'  
        run.font.size = Pt(15)  
        p.paragraph_format.first_line_indent = Pt(24)
        p.paragraph_format.line_spacing = Pt(14)
        p.alignment= WD_ALIGN_PARAGRAPH.CENTER

def add_page_numbers(document):
    sections = document.sections
    for section in sections:
        footer = section.footer
        p = footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        field = OxmlElement('w:fldSimple')
        field.set(qn('w:instr'), 'PAGE')
        run._element.append(field)

def main():
    document = Document()
    page_formatting(document)
    add_content(document)
    add_page_numbers(document)

    document.save("readersEdition.docx")
    subprocess.call(['powershell.exe', 'Start-Process', 'readersEdition.docx'])

main()
